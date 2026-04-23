"""
ARC Platform — Export Service
PDF and DOCX report generation from markdown report content.
"""
from __future__ import annotations

import io
import logging
import re
from datetime import datetime, timezone
from typing import Any

logger = logging.getLogger("arc.services.export")


class ExportService:
    """Generates PDF and DOCX exports from report data."""

    @staticmethod
    def generate_pdf(
        report: dict[str, Any],
        citations: list[dict[str, Any]],
        task: dict[str, Any],
    ) -> io.BytesIO:
        """
        Generate a formatted PDF report.

        Args:
            report: MongoDB report document with title, content, sections, etc.
            citations: List of citation documents.
            task: MongoDB task document.

        Returns:
            BytesIO stream containing the PDF.
        """
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch, mm
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            PageBreak,
            HRFlowable,
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=25 * mm,
            leftMargin=25 * mm,
            topMargin=30 * mm,
            bottomMargin=25 * mm,
        )

        styles = getSampleStyleSheet()

        # ── Custom styles ────────────────────────────────────────────────
        styles.add(ParagraphStyle(
            name="ReportTitle",
            parent=styles["Title"],
            fontSize=28,
            leading=34,
            spaceAfter=6,
            textColor=colors.HexColor("#1a1a2e"),
            alignment=TA_CENTER,
        ))
        styles.add(ParagraphStyle(
            name="ReportSubtitle",
            parent=styles["Normal"],
            fontSize=12,
            leading=16,
            spaceAfter=30,
            textColor=colors.HexColor("#6366f1"),
            alignment=TA_CENTER,
        ))
        styles.add(ParagraphStyle(
            name="SectionHeading",
            parent=styles["Heading1"],
            fontSize=18,
            leading=24,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor("#1e293b"),
            borderWidth=0,
            borderPadding=0,
        ))
        styles.add(ParagraphStyle(
            name="SubHeading",
            parent=styles["Heading2"],
            fontSize=14,
            leading=18,
            spaceBefore=14,
            spaceAfter=8,
            textColor=colors.HexColor("#334155"),
        ))
        styles.add(ParagraphStyle(
            name="BodyText2",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
            textColor=colors.HexColor("#374151"),
        ))
        styles.add(ParagraphStyle(
            name="BulletItem",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            leftIndent=20,
            spaceAfter=4,
            bulletIndent=8,
            textColor=colors.HexColor("#374151"),
        ))
        styles.add(ParagraphStyle(
            name="CodeBlock",
            parent=styles["Code"],
            fontSize=9,
            leading=12,
            spaceAfter=12,
            spaceBefore=8,
            backColor=colors.HexColor("#f1f5f9"),
            borderWidth=1,
            borderColor=colors.HexColor("#e2e8f0"),
            borderPadding=8,
            textColor=colors.HexColor("#1e293b"),
        ))
        styles.add(ParagraphStyle(
            name="CitationText",
            parent=styles["Normal"],
            fontSize=9,
            leading=13,
            spaceAfter=6,
            textColor=colors.HexColor("#64748b"),
        ))
        styles.add(ParagraphStyle(
            name="Footer",
            parent=styles["Normal"],
            fontSize=8,
            textColor=colors.HexColor("#94a3b8"),
            alignment=TA_CENTER,
        ))

        story: list[Any] = []

        # ── Title Page ───────────────────────────────────────────────────
        story.append(Spacer(1, 80))

        # Brand header
        brand_style = ParagraphStyle(
            name="Brand",
            parent=styles["Normal"],
            fontSize=14,
            textColor=colors.HexColor("#6366f1"),
            alignment=TA_CENTER,
            spaceAfter=8,
        )
        story.append(Paragraph("◆ ARC PLATFORM ◆", brand_style))
        story.append(Paragraph("Agentic Research Copilot", ParagraphStyle(
            name="BrandSub",
            parent=styles["Normal"],
            fontSize=10,
            textColor=colors.HexColor("#94a3b8"),
            alignment=TA_CENTER,
            spaceAfter=40,
        )))

        story.append(HRFlowable(
            width="60%", thickness=2, color=colors.HexColor("#6366f1"),
            spaceAfter=30, spaceBefore=0, hAlign="CENTER",
        ))

        title = _escape_html(report.get("title", task.get("title", "Research Report")))
        story.append(Paragraph(title, styles["ReportTitle"]))

        now = datetime.now(timezone.utc).strftime("%B %d, %Y")
        story.append(Paragraph(f"Generated on {now}", styles["ReportSubtitle"]))

        story.append(HRFlowable(
            width="60%", thickness=1, color=colors.HexColor("#e2e8f0"),
            spaceAfter=20, spaceBefore=20, hAlign="CENTER",
        ))

        # Summary on title page
        summary = report.get("summary", "")
        if summary:
            story.append(Spacer(1, 20))
            story.append(Paragraph("<b>Executive Summary</b>", ParagraphStyle(
                name="SumTitle",
                parent=styles["Normal"],
                fontSize=12,
                textColor=colors.HexColor("#1e293b"),
                alignment=TA_CENTER,
                spaceAfter=10,
            )))
            for para in summary.split("\n"):
                para = para.strip()
                if para:
                    story.append(Paragraph(_escape_html(para), ParagraphStyle(
                        name="SumBody",
                        parent=styles["Normal"],
                        fontSize=10,
                        leading=15,
                        textColor=colors.HexColor("#475569"),
                        alignment=TA_JUSTIFY,
                        spaceAfter=6,
                        leftIndent=30,
                        rightIndent=30,
                    )))

        story.append(PageBreak())

        # ── Table of Contents ────────────────────────────────────────────
        sections = report.get("sections", [])
        if sections:
            story.append(Paragraph("Table of Contents", styles["SectionHeading"]))
            story.append(HRFlowable(
                width="100%", thickness=1, color=colors.HexColor("#e2e8f0"),
                spaceAfter=16,
            ))
            for i, sec in enumerate(sections, 1):
                sec_title = _escape_html(sec.get("title", f"Section {i}"))
                toc_style = ParagraphStyle(
                    name=f"TOC_{i}",
                    parent=styles["Normal"],
                    fontSize=11,
                    leading=20,
                    textColor=colors.HexColor("#334155"),
                    leftIndent=10,
                )
                story.append(Paragraph(f"{i}. {sec_title}", toc_style))
            story.append(Spacer(1, 30))
            story.append(HRFlowable(
                width="100%", thickness=1, color=colors.HexColor("#e2e8f0"),
                spaceAfter=20,
            ))

        # ── Report Body ──────────────────────────────────────────────────
        content = report.get("content", "")
        if content:
            _add_markdown_content(story, content, styles)
        elif sections:
            for i, sec in enumerate(sections, 1):
                sec_title = _escape_html(sec.get("title", f"Section {i}"))
                story.append(Paragraph(f"{i}. {sec_title}", styles["SectionHeading"]))
                sec_content = sec.get("content", "")
                if sec_content:
                    _add_markdown_content(story, sec_content, styles)
                story.append(Spacer(1, 10))

        # ── Citation Appendix ────────────────────────────────────────────
        if citations:
            story.append(PageBreak())
            story.append(Paragraph("References & Citations", styles["SectionHeading"]))
            story.append(HRFlowable(
                width="100%", thickness=1, color=colors.HexColor("#e2e8f0"),
                spaceAfter=16,
            ))

            for i, cit in enumerate(citations, 1):
                cit_title = _escape_html(cit.get("title", "Unknown"))
                cit_url = cit.get("url", "")
                cit_type = cit.get("citation_type", "web")
                cit_excerpt = _escape_html(cit.get("excerpt", "")) if cit.get("excerpt") else ""

                entry = f"<b>[{i}]</b> {cit_title}"
                if cit_url:
                    entry += f' — <font color="#6366f1"><link href="{cit_url}">{cit_url}</link></font>'
                entry += f' <font color="#94a3b8">[{cit_type}]</font>'
                if cit_excerpt:
                    entry += f'<br/><font color="#64748b"><i>"{cit_excerpt[:200]}"</i></font>'

                story.append(Paragraph(entry, styles["CitationText"]))

        # ── Build PDF ────────────────────────────────────────────────────
        def _page_footer(canvas, doc_obj):
            canvas.saveState()
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(colors.HexColor("#94a3b8"))
            canvas.drawCentredString(
                A4[0] / 2, 15 * mm,
                f"ARC Platform Report — Page {canvas.getPageNumber()}"
            )
            canvas.restoreState()

        doc.build(story, onFirstPage=_page_footer, onLaterPages=_page_footer)
        buffer.seek(0)

        logger.info("PDF generated for task %s (%d bytes)", task.get("_id", "?"), buffer.getbuffer().nbytes)
        return buffer

    @staticmethod
    def generate_docx(
        report: dict[str, Any],
        citations: list[dict[str, Any]],
        task: dict[str, Any],
    ) -> io.BytesIO:
        """
        Generate a formatted DOCX report.

        Args:
            report: MongoDB report document.
            citations: List of citation documents.
            task: MongoDB task document.

        Returns:
            BytesIO stream containing the DOCX.
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Document styles ──────────────────────────────────────────────
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(55, 65, 81)

        # ── Title Page ───────────────────────────────────────────────────
        # Brand
        brand = doc.add_paragraph()
        brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = brand.add_run("◆ ARC PLATFORM ◆")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(99, 102, 241)
        run.bold = True

        brand_sub = doc.add_paragraph()
        brand_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = brand_sub.add_run("Agentic Research Copilot")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(148, 163, 184)

        doc.add_paragraph()  # Spacer

        # Title
        title_text = report.get("title", task.get("title", "Research Report"))
        title_para = doc.add_heading(title_text, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(26, 26, 46)

        # Date
        now = datetime.now(timezone.utc).strftime("%B %d, %Y")
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Generated on {now}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(99, 102, 241)

        # Summary
        summary = report.get("summary", "")
        if summary:
            doc.add_paragraph()
            heading = doc.add_heading("Executive Summary", level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(30, 41, 59)
            for para_text in summary.split("\n"):
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

        # ── Table of Contents Placeholder ────────────────────────────────
        sections = report.get("sections", [])
        if sections:
            heading = doc.add_heading("Table of Contents", level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(30, 41, 59)
            for i, sec in enumerate(sections, 1):
                sec_title = sec.get("title", f"Section {i}")
                toc_p = doc.add_paragraph(f"{i}. {sec_title}")
                toc_p.paragraph_format.left_indent = Inches(0.3)
                toc_p.paragraph_format.space_after = Pt(4)
            doc.add_paragraph()

        # ── Report Body ──────────────────────────────────────────────────
        content = report.get("content", "")
        if content:
            _add_docx_markdown_content(doc, content)
        elif sections:
            for i, sec in enumerate(sections, 1):
                sec_title = sec.get("title", f"Section {i}")
                heading = doc.add_heading(f"{i}. {sec_title}", level=1)
                for run_obj in heading.runs:
                    run_obj.font.color.rgb = RGBColor(30, 41, 59)
                sec_content = sec.get("content", "")
                if sec_content:
                    _add_docx_markdown_content(doc, sec_content)

        # ── Citation Appendix ────────────────────────────────────────────
        if citations:
            doc.add_page_break()
            heading = doc.add_heading("References & Citations", level=1)
            for run_obj in heading.runs:
                run_obj.font.color.rgb = RGBColor(30, 41, 59)

            for i, cit in enumerate(citations, 1):
                cit_title = cit.get("title", "Unknown")
                cit_url = cit.get("url", "")
                cit_type = cit.get("citation_type", "web")
                cit_excerpt = cit.get("excerpt", "")

                p = doc.add_paragraph()
                run = p.add_run(f"[{i}] ")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 41, 59)

                run = p.add_run(cit_title)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 65, 85)

                if cit_url:
                    run = p.add_run(f" — {cit_url}")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(99, 102, 241)

                run = p.add_run(f" [{cit_type}]")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(148, 163, 184)

                if cit_excerpt:
                    excerpt_p = doc.add_paragraph()
                    run = excerpt_p.add_run(f'"{cit_excerpt[:200]}"')
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(100, 116, 139)
                    excerpt_p.paragraph_format.left_indent = Inches(0.5)

                p.paragraph_format.space_after = Pt(4)

        # ── Save ─────────────────────────────────────────────────────────
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logger.info("DOCX generated for task %s (%d bytes)", task.get("_id", "?"), buffer.getbuffer().nbytes)
        return buffer


# ── Helper functions ─────────────────────────────────────────────────────────
def _escape_html(text: str) -> str:
    """Escape HTML special characters for ReportLab Paragraph."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _add_markdown_content(story: list, content: str, styles: Any) -> None:
    """Parse markdown content into ReportLab flowables."""
    lines = content.split("\n")
    in_code_block = False
    code_buffer: list[str] = []

    for line in lines:
        # Code block boundaries
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                code_text = _escape_html("\n".join(code_buffer))
                story.append(Paragraph(
                    code_text.replace("\n", "<br/>"),
                    styles["CodeBlock"],
                ))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        stripped = line.strip()

        # Empty line
        if not stripped:
            from reportlab.platypus import Spacer
            story.append(Spacer(1, 6))
            continue

        # Headings
        if stripped.startswith("#### "):
            text = _escape_html(stripped[5:])
            story.append(Paragraph(text, styles["SubHeading"]))
        elif stripped.startswith("### "):
            text = _escape_html(stripped[4:])
            story.append(Paragraph(text, styles["SubHeading"]))
        elif stripped.startswith("## "):
            text = _escape_html(stripped[3:])
            story.append(Paragraph(text, styles["SectionHeading"]))
        elif stripped.startswith("# "):
            text = _escape_html(stripped[2:])
            story.append(Paragraph(text, styles["SectionHeading"]))
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = _escape_html(stripped[2:])
            text = _apply_inline_formatting(text)
            story.append(Paragraph(f"• {text}", styles["BulletItem"]))
        elif re.match(r"^\d+\.\s", stripped):
            text = _escape_html(re.sub(r"^\d+\.\s", "", stripped))
            text = _apply_inline_formatting(text)
            num = re.match(r"^(\d+)\.", stripped)
            prefix = num.group(1) if num else "•"
            story.append(Paragraph(f"{prefix}. {text}", styles["BulletItem"]))
        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            from reportlab.lib import colors as _c
            from reportlab.platypus import HRFlowable
            story.append(HRFlowable(
                width="100%", thickness=1, color=_c.HexColor("#e2e8f0"),
                spaceAfter=10, spaceBefore=10,
            ))
        # Regular paragraph
        else:
            text = _escape_html(stripped)
            text = _apply_inline_formatting(text)
            story.append(Paragraph(text, styles["BodyText2"]))


def _apply_inline_formatting(text: str) -> str:
    """Convert markdown bold/italic to ReportLab XML tags."""
    # Bold
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # Italic
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    # Inline code
    text = re.sub(r"`(.+?)`", r'<font face="Courier" size="9" color="#1e293b">\1</font>', text)
    return text


def _add_docx_markdown_content(doc: Any, content: str) -> None:
    """Parse markdown content into python-docx paragraphs."""
    from docx.shared import Pt, RGBColor

    lines = content.split("\n")
    in_code_block = False
    code_buffer: list[str] = []

    for line in lines:
        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 41, 59)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith("#### "):
            h = doc.add_heading(stripped[5:], level=4)
            for r in h.runs:
                r.font.color.rgb = RGBColor(51, 65, 85)
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            for r in h.runs:
                r.font.color.rgb = RGBColor(51, 65, 85)
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(30, 41, 59)
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(30, 41, 59)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(text, style="List Number")
            p.paragraph_format.space_after = Pt(4)
        # Horizontal rule — skip
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 50)
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_docx_inline_formatting(p, stripped)
            p.paragraph_format.space_after = Pt(6)


def _add_docx_inline_formatting(paragraph: Any, text: str) -> None:
    """Add text with inline bold/italic to a docx paragraph."""
    from docx.shared import Pt, RGBColor

    # Split on bold/italic markers
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(30, 41, 59)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(55, 65, 81)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(30, 41, 59)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(55, 65, 81)
