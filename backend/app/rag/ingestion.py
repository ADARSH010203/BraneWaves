"""
ARC Platform — Document Ingestion
Extracts text from uploaded files for RAG processing.
"""
from __future__ import annotations

import io
import logging
from typing import Any

logger = logging.getLogger("arc.rag.ingestion")

# Supported content types
SUPPORTED_TYPES = {
    "text/plain": "txt",
    "text/markdown": "md",
    "text/csv": "csv",
    "application/json": "json",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
}


async def extract_text(file_bytes: bytes, content_type: str, filename: str) -> str:
    """
    Extract text content from file bytes based on content type.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Plain text formats
    if content_type in ("text/plain", "text/markdown", "text/csv") or ext in ("txt", "md", "csv"):
        return file_bytes.decode("utf-8", errors="replace")

    # JSON
    if content_type == "application/json" or ext == "json":
        import json
        try:
            data = json.loads(file_bytes.decode("utf-8"))
            return json.dumps(data, indent=2)
        except json.JSONDecodeError:
            return file_bytes.decode("utf-8", errors="replace")

    # PDF
    if content_type == "application/pdf" or ext == "pdf":
        return _extract_pdf(file_bytes)

    # DOCX / DOC
    if content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or ext in ("docx", "doc"):
        return _extract_docx(file_bytes)

    # Fallback: try as text
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception:
        logger.warning("Could not extract text from %s (%s)", filename, content_type)
        return ""


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyPDF2 or pdfplumber."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
    except ImportError:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except ImportError:
        logger.warning("No PDF library available. Install pdfplumber or PyPDF2.")
        return "[PDF content - extraction library not available]"


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    table_rows.append(row_text)
        return "\n\n".join(paragraphs) + ("\n\n" + "\n".join(table_rows) if table_rows else "")
    except ImportError:
        logger.warning("python-docx not installed")
        return "[DOCX - install python-docx]"
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        return ""
